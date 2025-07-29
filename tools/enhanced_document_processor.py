"""
Enhanced Document Processor with Advanced Extraction
Supports PDF (text + tables + OCR), Word (multiple formats), and Excel optimization
"""

import os
import logging
import pandas as pd
import json
import hashlib
import re
from typing import List, Dict, Any, Optional, Tuple, Union
from datetime import datetime
import traceback

# Enhanced PDF processing imports
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
    print("✅ DEBUG: PyMuPDF (fitz) available")
except ImportError as e:
    PYMUPDF_AVAILABLE = False
    print(f"⚠️ DEBUG: PyMuPDF (fitz) not available: {e}")

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

# OCR support
try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

# Enhanced Word processing
try:
    import mammoth
    import docx
    WORD_ENHANCED = True
except ImportError:
    WORD_ENHANCED = False

logger = logging.getLogger(__name__)

class EnhancedDocumentProcessor:
    """Advanced document processor with intelligent extraction methods"""
    
    def __init__(self, folder_path: str):
        self.folder_path = folder_path
        self.supported_extensions = {'.pdf', '.xlsx', '.xls', '.txt', '.docx', '.doc'}
        
        # Ensure folder exists
        if not os.path.exists(folder_path):
            os.makedirs(folder_path, exist_ok=True)
            print(f"📁 Created folder: {folder_path}")
        
        # Initialize extraction statistics
        self.extraction_stats = {
            'total_files': 0,
            'successful_extractions': 0,
            'failed_extractions': 0,
            'methods_used': {}
        }
    
    def extract_from_pdf_advanced(self, file_path: str) -> Dict[str, Any]:
        """Advanced PDF extraction with multiple methods and quality scoring"""
        print(f"🔍 Advanced PDF extraction: {os.path.basename(file_path)}")
        
        extraction_results = {
            'methods': {},
            'best_text': '',
            'best_method': '',
            'tables_found': [],
            'images_processed': 0,
            'quality_score': 0
        }
        
        # Method 1: PyMuPDF (best for general text)
        if PYMUPDF_AVAILABLE:
            try:
                text = self._extract_with_pymupdf_advanced(file_path)
                extraction_results['methods']['pymupdf'] = {
                    'text': text,
                    'length': len(text),
                    'quality': self._calculate_text_quality(text)
                }
                print(f"✅ PyMuPDF: {len(text)} chars, quality: {extraction_results['methods']['pymupdf']['quality']:.2f}")
            except Exception as e:
                print(f"❌ PyMuPDF failed: {str(e)}")
        
        # Method 2: pdfplumber (excellent for tables and structured data)
        if PDFPLUMBER_AVAILABLE:
            try:
                result = self._extract_with_pdfplumber(file_path)
                extraction_results['methods']['pdfplumber'] = {
                    'text': result['text'],
                    'length': len(result['text']),
                    'quality': self._calculate_text_quality(result['text']),
                    'tables': result['tables']
                }
                extraction_results['tables_found'] = result['tables']
                print(f"✅ pdfplumber: {len(result['text'])} chars, {len(result['tables'])} tables, quality: {extraction_results['methods']['pdfplumber']['quality']:.2f}")
            except Exception as e:
                print(f"❌ pdfplumber failed: {str(e)}")
        
        # Method 3: PyPDF2 (fallback)
        if PYPDF2_AVAILABLE:
            try:
                text = self._extract_with_pypdf2_enhanced(file_path)
                extraction_results['methods']['pypdf2'] = {
                    'text': text,
                    'length': len(text),
                    'quality': self._calculate_text_quality(text)
                }
                print(f"✅ PyPDF2: {len(text)} chars, quality: {extraction_results['methods']['pypdf2']['quality']:.2f}")
            except Exception as e:
                print(f"❌ PyPDF2 failed: {str(e)}")
        
        # Method 4: OCR for scanned PDFs
        if OCR_AVAILABLE and self._is_scanned_pdf(file_path):
            try:
                text = self._extract_with_ocr(file_path)
                extraction_results['methods']['ocr'] = {
                    'text': text,
                    'length': len(text),
                    'quality': self._calculate_text_quality(text)
                }
                print(f"✅ OCR: {len(text)} chars, quality: {extraction_results['methods']['ocr']['quality']:.2f}")
            except Exception as e:
                print(f"❌ OCR failed: {str(e)}")
        
        # Select best method
        best_method, best_text = self._select_best_extraction(extraction_results['methods'])
        extraction_results['best_method'] = best_method
        extraction_results['best_text'] = best_text
        extraction_results['quality_score'] = extraction_results['methods'].get(best_method, {}).get('quality', 0)
        
        print(f"🏆 Best method: {best_method} with quality score: {extraction_results['quality_score']:.2f}")
        
        return extraction_results
    
    def _extract_with_pymupdf_advanced(self, file_path: str) -> str:
        """Enhanced PyMuPDF extraction with better text handling"""
        doc = fitz.open(file_path)
        text_parts = []
        
        for page_num in range(doc.page_count):
            page = doc[page_num]
            
            # Try different extraction methods for each page
            methods = {
                'standard': page.get_text(),
                'dict': self._extract_text_from_dict(page.get_text("dict")),
                'blocks': self._extract_text_from_blocks(page.get_text("blocks"))
            }
            
            # Choose best method for this page
            best_page_text = max(methods.values(), key=len)
            if best_page_text.strip():
                text_parts.append(f"--- Página {page_num + 1} ---")
                text_parts.append(best_page_text)
        
        doc.close()
        return "\n".join(text_parts)
    
    def _extract_with_pdfplumber(self, file_path: str) -> Dict[str, Any]:
        """Extract text and tables using pdfplumber"""
        text_parts = []
        tables = []
        
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                # Extract text
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"--- Página {page_num + 1} ---")
                    text_parts.append(page_text)
                
                # Extract tables
                page_tables = page.extract_tables()
                for table_idx, table in enumerate(page_tables):
                    if table:
                        # Convert table to readable format
                        table_text = self._format_table_as_text(table, page_num + 1, table_idx + 1)
                        tables.append(table_text)
                        text_parts.append(f"\n--- TABLA {page_num + 1}.{table_idx + 1} ---")
                        text_parts.append(table_text)
        
        return {
            'text': "\n".join(text_parts),
            'tables': tables
        }
    
    def _extract_with_pypdf2_enhanced(self, file_path: str) -> str:
        """Enhanced PyPDF2 extraction with better error handling"""
        text_parts = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(f"--- Página {page_num + 1} ---")
                        text_parts.append(page_text)
                except Exception as e:
                    print(f"⚠️ Error extracting page {page_num + 1}: {str(e)}")
        
        return "\n".join(text_parts)
    
    def _extract_with_ocr(self, file_path: str) -> str:
        """Extract text using OCR for scanned PDFs"""
        text_parts = []
        
        doc = fitz.open(file_path)
        for page_num in range(doc.page_count):
            page = doc[page_num]
            
            # Convert page to image
            pix = page.get_pixmap()
            img_data = pix.tobytes("ppm")
            
            # Convert to PIL Image
            from io import BytesIO
            img = Image.open(BytesIO(img_data))
            
            # Extract text with OCR
            ocr_text = pytesseract.image_to_string(img, lang='spa+eng')
            
            if ocr_text.strip():
                text_parts.append(f"--- Página {page_num + 1} (OCR) ---")
                text_parts.append(ocr_text)
        
        doc.close()
        return "\n".join(text_parts)
    
    def extract_from_word_advanced(self, file_path: str) -> Dict[str, Any]:
        """Enhanced Word document extraction"""
        print(f"📄 Advanced Word extraction: {os.path.basename(file_path)}")
        
        extraction_results = {
            'methods': {},
            'best_text': '',
            'best_method': '',
            'quality_score': 0
        }
        
        # Method 1: mammoth (best for formatting preservation)
        if WORD_ENHANCED:
            try:
                with open(file_path, "rb") as docx_file:
                    result = mammoth.extract_raw_text(docx_file)
                    text = result.value
                    
                extraction_results['methods']['mammoth'] = {
                    'text': text,
                    'length': len(text),
                    'quality': self._calculate_text_quality(text)
                }
                print(f"✅ mammoth: {len(text)} chars, quality: {extraction_results['methods']['mammoth']['quality']:.2f}")
            except Exception as e:
                print(f"❌ mammoth failed: {str(e)}")
        
        # Method 2: python-docx (structured approach)
        if WORD_ENHANCED:
            try:
                doc = docx.Document(file_path)
                paragraphs = []
                
                for para in doc.paragraphs:
                    if para.text.strip():
                        paragraphs.append(para.text)
                
                # Extract tables
                for table in doc.tables:
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text.strip() for cell in row.cells]
                        table_data.append(" | ".join(row_data))
                    paragraphs.append("\n--- TABLA ---")
                    paragraphs.extend(table_data)
                
                text = "\n".join(paragraphs)
                
                extraction_results['methods']['python_docx'] = {
                    'text': text,
                    'length': len(text),
                    'quality': self._calculate_text_quality(text)
                }
                print(f"✅ python-docx: {len(text)} chars, quality: {extraction_results['methods']['python_docx']['quality']:.2f}")
            except Exception as e:
                print(f"❌ python-docx failed: {str(e)}")
        
        # Select best method
        best_method, best_text = self._select_best_extraction(extraction_results['methods'])
        extraction_results['best_method'] = best_method
        extraction_results['best_text'] = best_text
        extraction_results['quality_score'] = extraction_results['methods'].get(best_method, {}).get('quality', 0)
        
        print(f"🏆 Best Word method: {best_method} with quality score: {extraction_results['quality_score']:.2f}")
        
        return extraction_results
    
    def _calculate_text_quality(self, text: str) -> float:
        """Calculate text quality score based on various factors"""
        if not text or not text.strip():
            return 0.0
        
        score = 0.0
        
        # Length score (more text usually better, but with diminishing returns)
        length_score = min(len(text) / 1000, 1.0) * 0.3
        score += length_score
        
        # Word count score
        words = text.split()
        word_count_score = min(len(words) / 200, 1.0) * 0.2
        score += word_count_score
        
        # Character variety (avoid repetitive extraction errors)
        unique_chars = len(set(text.lower()))
        variety_score = min(unique_chars / 50, 1.0) * 0.2
        score += variety_score
        
        # Sentence structure score
        sentences = text.count('.') + text.count('!') + text.count('?')
        sentence_score = min(sentences / 50, 1.0) * 0.1
        score += sentence_score
        
        # Reduce score for excessive whitespace or repetitive patterns
        whitespace_ratio = text.count(' ') / len(text) if len(text) > 0 else 0
        if whitespace_ratio > 0.3:  # Too much whitespace
            score *= 0.8
        
        # Reduce score for excessive newlines
        newline_ratio = text.count('\n') / len(text) if len(text) > 0 else 0
        if newline_ratio > 0.1:  # Too many newlines
            score *= 0.9
        
        # Bonus for structured content (headings, lists, etc.)
        if any(marker in text for marker in ['•', '-', '1.', '2.', '3.', '###', '---']):
            score *= 1.1
        
        return min(score, 1.0)  # Cap at 1.0
    
    def _select_best_extraction(self, methods: Dict[str, Dict]) -> Tuple[str, str]:
        """Select the best extraction method based on quality scores"""
        if not methods:
            return 'none', ''
        
        best_method = ''
        best_score = 0
        best_text = ''
        
        for method, data in methods.items():
            quality = data.get('quality', 0)
            length = data.get('length', 0)
            
            # Combine quality and length scores
            combined_score = quality * 0.7 + min(length / 2000, 1.0) * 0.3
            
            if combined_score > best_score:
                best_score = combined_score
                best_method = method
                best_text = data.get('text', '')
        
        return best_method, best_text
    
    def _is_scanned_pdf(self, file_path: str) -> bool:
        """Check if PDF is likely scanned (image-based)"""
        if not PYMUPDF_AVAILABLE:
            return False
        
        try:
            doc = fitz.open(file_path)
            text_pages = 0
            total_pages = doc.page_count
            
            # Check first few pages
            check_pages = min(3, total_pages)
            
            for page_num in range(check_pages):
                page = doc[page_num]
                text = page.get_text().strip()
                if len(text) > 50:  # Has meaningful text
                    text_pages += 1
            
            doc.close()
            
            # If less than half of checked pages have text, likely scanned
            return text_pages / check_pages < 0.5
        except:
            return False
    
    def _format_table_as_text(self, table: List[List[str]], page_num: int, table_num: int) -> str:
        """Format extracted table as readable text"""
        if not table:
            return ""
        
        lines = [f"TABLA {page_num}.{table_num}:"]
        
        for row_idx, row in enumerate(table):
            if row and any(cell for cell in row if cell and cell.strip()):
                # Clean cells
                clean_cells = [str(cell).strip() if cell else "" for cell in row]
                # Join with separator
                line = " | ".join(clean_cells)
                lines.append(line)
                
                # Add separator after header
                if row_idx == 0 and len(table) > 1:
                    lines.append("-" * len(line))
        
        return "\n".join(lines)
    
    def _extract_text_from_dict(self, text_dict: Dict) -> str:
        """Extract text from PyMuPDF dictionary structure"""
        text_parts = []
        
        for block in text_dict.get("blocks", []):
            if "lines" in block:
                for line in block["lines"]:
                    line_text = ""
                    for span in line.get("spans", []):
                        line_text += span.get("text", "")
                    if line_text.strip():
                        text_parts.append(line_text.strip())
        
        return "\n".join(text_parts)
    
    def _extract_text_from_blocks(self, blocks: List) -> str:
        """Extract text from PyMuPDF blocks"""
        text_parts = []
        
        for block in blocks:
            if len(block) >= 5:  # Text block
                text = block[4].strip()
                if text:
                    text_parts.append(text)
        
        return "\n".join(text_parts)
    
    def process_document_intelligent(self, file_path: str) -> Dict[str, Any]:
        """Process document with intelligent method selection"""
        _, ext = os.path.splitext(file_path.lower())
        
        result = {
            'file_path': file_path,
            'file_type': ext[1:] if ext else 'unknown',
            'extraction_method': 'none',
            'text': '',
            'metadata': {},
            'success': False
        }
        
        try:
            if ext == '.pdf':
                extraction_result = self.extract_from_pdf_advanced(file_path)
                result['text'] = extraction_result['best_text']
                result['extraction_method'] = extraction_result['best_method']
                result['metadata'] = {
                    'quality_score': extraction_result['quality_score'],
                    'tables_found': len(extraction_result['tables_found']),
                    'methods_tried': list(extraction_result['methods'].keys())
                }
                result['success'] = bool(result['text'].strip())
            
            elif ext in ['.docx', '.doc']:
                extraction_result = self.extract_from_word_advanced(file_path)
                result['text'] = extraction_result['best_text']
                result['extraction_method'] = extraction_result['best_method']
                result['metadata'] = {
                    'quality_score': extraction_result['quality_score'],
                    'methods_tried': list(extraction_result['methods'].keys())
                }
                result['success'] = bool(result['text'].strip())
            
            elif ext in ['.xlsx', '.xls']:
                # Excel processing (existing logic can be enhanced here)
                result['text'] = self._extract_from_excel(file_path)
                result['extraction_method'] = 'pandas'
                result['success'] = bool(result['text'].strip())
            
            elif ext == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    result['text'] = f.read()
                result['extraction_method'] = 'text_file'
                result['success'] = True
        
        except Exception as e:
            result['error'] = str(e)
            print(f"❌ Error processing {file_path}: {str(e)}")
        
        # Update statistics
        self.extraction_stats['total_files'] += 1
        if result['success']:
            self.extraction_stats['successful_extractions'] += 1
            method = result['extraction_method']
            self.extraction_stats['methods_used'][method] = self.extraction_stats['methods_used'].get(method, 0) + 1
        else:
            self.extraction_stats['failed_extractions'] += 1
        
        return result
    
    def _extract_from_excel(self, file_path: str) -> str:
        """Enhanced Excel extraction"""
        text_parts = []
        
        # Read all sheets
        excel_file = pd.ExcelFile(file_path)
        
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            
            if not df.empty:
                text_parts.append(f"--- HOJA: {sheet_name} ---")
                
                # Convert DataFrame to readable text
                for index, row in df.iterrows():
                    row_text = " | ".join([str(val) for val in row.values if pd.notna(val)])
                    if row_text.strip():
                        text_parts.append(row_text)
        
        return "\n".join(text_parts)
    
    def get_extraction_report(self) -> Dict[str, Any]:
        """Get detailed extraction statistics"""
        return {
            'statistics': self.extraction_stats.copy(),
            'success_rate': (self.extraction_stats['successful_extractions'] / 
                           max(self.extraction_stats['total_files'], 1)) * 100,
            'most_used_method': max(self.extraction_stats['methods_used'].items(), 
                                  key=lambda x: x[1])[0] if self.extraction_stats['methods_used'] else 'none'
        }

def test_enhanced_processor():
    """Test the enhanced document processor"""
    processor = EnhancedDocumentProcessor("folders/exhibitors")
    
    # Test with the catalog PDF
    catalog_path = "folders/exhibitors/Catálogo Mobiliario Espacio Food & Service 2025.pdf"
    
    if os.path.exists(catalog_path):
        print("🧪 Testing enhanced processor with furniture catalog...")
        result = processor.process_document_intelligent(catalog_path)
        
        print(f"\n📊 EXTRACTION RESULTS:")
        print(f"Success: {result['success']}")
        print(f"Method: {result['extraction_method']}")
        print(f"Text length: {len(result['text'])} characters")
        print(f"Quality score: {result['metadata'].get('quality_score', 'N/A')}")
        print(f"Tables found: {result['metadata'].get('tables_found', 0)}")
        
        print(f"\n📝 SAMPLE TEXT (first 500 chars):")
        print("-" * 50)
        print(result['text'][:500] + "..." if len(result['text']) > 500 else result['text'])
        
        print(f"\n📈 EXTRACTION REPORT:")
        report = processor.get_extraction_report()
        for key, value in report.items():
            print(f"{key}: {value}")
        
        return result
    else:
        print(f"❌ Catalog file not found: {catalog_path}")
        return None

if __name__ == "__main__":
    test_enhanced_processor()